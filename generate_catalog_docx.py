"""
Gera o DOCX do catalogo Cyber Informatica — TABELA DE PRECOS individualizada.
Iago pode editar livremente no Word/LibreOffice.

Uso: python generate_catalog_docx.py
Saida: assets/catalogo-cyber.docx

ATUALIZADO 2026-06-26 (versao DOCX editavel, foco em lojista)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

# ============================================================================
# Paleta Cyber (em hex para usar em cores RGB)
# ============================================================================
CYBER_BLUE = '0066FF'
CYBER_BLUE_HOVER = '0052CC'
CIRCUIT_GREEN = '00CC6A'
NAVY_950 = '050A14'
NAVY_900 = '0A1929'
GRAY_700 = '374151'
GRAY_500 = '6B7280'
GRAY_200 = 'E5E7EB'
WHITE = 'FFFFFF'
LIGHT_BG = 'F5F8FF'

OUTPUT = os.path.join(
    os.path.dirname(os.path.abspath(__file__)),
    'assets', 'catalogo-cyber.docx'
)

# ============================================================================
# Tabela de precos — espelha lib/pricing-models.ts
# Formato: (modelo, valor_lojista_35%, valor_cheio_referencia)
# Para editar: alterar os valores e rodar 'python generate_catalog_docx.py'
# ============================================================================

GROUPS = [
    ('Apple', [
        ('iPhone 17 Pro Max',   1500, 4000),
        ('iPhone 17 Pro',       1250, 3600),
        ('iPhone 17 Air',       1100, 3200),
        ('iPhone 17',           1000, 2800),
        ('iPhone 16 Pro Max',   1200, 3500),
        ('iPhone 16 Pro',       1100, 3200),
        ('iPhone 16 Plus',      1000, 2800),
        ('iPhone 16',            900, 2500),
        ('iPhone 15 Pro Max',   1050, 3000),
        ('iPhone 15 Pro',       1000, 2800),
        ('iPhone 15 Plus',       900, 2500),
        ('iPhone 15',            750, 2200),
        ('iPhone 14 Pro Max',   1000, 2900),
        ('iPhone 14 Pro',        950, 2700),
        ('iPhone 14 Plus',       800, 2300),
        ('iPhone 14',            700, 2000),
        ('iPhone 13 Pro Max',    900, 2500),
        ('iPhone 13 Pro',        800, 2300),
        ('iPhone 13',            600, 1700),
        ('iPhone 12 Pro Max',    700, 2000),
        ('iPhone 12 Pro',        650, 1800),
        ('iPhone 12',            500, 1500),
        ('iPhone 11 Pro Max',    600, 1700),
        ('iPhone 11 Pro',        500, 1500),
        ('iPhone 11',            450, 1300),
        ('iPhone XS Max',        400, 1100),
        ('iPhone XS',            350, 1000),
        ('iPhone XR',            300, 900),
        ('iPhone SE (3a ger)',   300, 800),
    ]),
    ('Samsung Galaxy S', [
        ('Galaxy S25 Ultra',    1350, 3800),
        ('Galaxy S25+',         1100, 3200),
        ('Galaxy S25',          1000, 2800),
        ('Galaxy S24 Ultra',    1000, 2800),
        ('Galaxy S24+',          750, 2200),
        ('Galaxy S24',           650, 1800),
        ('Galaxy S24 FE',        600, 1700),
        ('Galaxy S23 Ultra',     900, 2500),
        ('Galaxy S23+',          700, 2000),
        ('Galaxy S23',           600, 1700),
        ('Galaxy S23 FE',        500, 1500),
        ('Galaxy S22 Ultra',     700, 2000),
        ('Galaxy S22+',          550, 1600),
        ('Galaxy S22',           500, 1500),
    ]),
    ('Samsung Galaxy A', [
        ('Galaxy A55',  400, 1200),
        ('Galaxy A54',  400, 1100),
        ('Galaxy A35',  300, 850),
        ('Galaxy A34',  300, 800),
        ('Galaxy A25',  250, 700),
        ('Galaxy A24',  250, 650),
        ('Galaxy A15',  250, 650),
        ('Galaxy A14',  200, 600),
        ('Galaxy A05',  175, 500),
    ]),
    ('Xiaomi', [
        ('Redmi Note 14 Pro',   400, 1100),
        ('Redmi Note 14',       250, 700),
        ('Redmi Note 13 Pro',   300, 900),
        ('Redmi Note 13',       250, 700),
        ('Redmi 14C',           200, 600),
        ('Redmi 13C',           200, 550),
        ('Poco X6 Pro',         450, 1300),
        ('Poco X6',             400, 1100),
        ('Poco X5 Pro',         300, 900),
    ]),
    ('Motorola', [
        ('Moto G84',     300, 900),
        ('Moto G54',     200, 600),
        ('Moto G34',     200, 550),
        ('Moto G24',     175, 500),
        ('Moto G14',     175, 480),
        ('Moto G04',     150, 450),
        ('Edge 50 Pro',  450, 1300),
        ('Edge 50',      400, 1100),
        ('Edge 40',      350, 1000),
    ]),
    ('Google Pixel', [
        ('Pixel 9 Pro XL',  950, 2700),
        ('Pixel 9 Pro',     850, 2400),
        ('Pixel 9',         650, 1800),
        ('Pixel 8 Pro',     800, 2300),
        ('Pixel 8',         500, 1500),
        ('Pixel 8a',        400, 1100),
    ]),
]

TOTAL_MODELS = sum(len(models) for _, models in GROUPS)


# ============================================================================
# Helpers de formatacao DOCX
# ============================================================================

def set_cell_bg(cell, hex_color: str):
    """Define cor de fundo da celula."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color='E5E7EB', size='4'):
    """Adiciona bordas finas em todas as direcoes."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), size)
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def brl(v: int) -> str:
    return f'R$ {v:,.0f}'.replace(',', '.')


def add_paragraph(doc, text='', size=11, bold=False, color=None,
                  align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.name = 'Calibri'
        if bold:
            run.bold = True
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    return p


def style_cell(cell, text, size=10, bold=False, color=None,
               align=WD_ALIGN_PARAGRAPH.LEFT, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = 'Calibri'
    if bold:
        run.bold = True
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if bg:
        set_cell_bg(cell, bg)
    # Remove espacos extras na celula
    cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)


def add_brand_table(doc, brand: str, models: list):
    """Adiciona cabecalho da marca + tabela de modelos."""
    # Cabecalho da marca
    th = doc.add_table(rows=1, cols=1)
    th.alignment = WD_TABLE_ALIGNMENT.CENTER
    th.autofit = False
    cell = th.rows[0].cells[0]
    cell.width = Cm(17)
    style_cell(cell, brand.upper(), size=10, bold=True,
               color=WHITE, bg=NAVY_950,
               align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_borders(cell, color=NAVY_950, size='4')

    # Tabela de modelos (3 colunas)
    t = doc.add_table(rows=len(models) + 1, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    # Larguras: 8cm / 4.5cm / 4.5cm
    widths = [Cm(8.5), Cm(4.25), Cm(4.25)]
    for i, w in enumerate(widths):
        for cell in t.columns[i].cells:
            cell.width = w

    # Header row
    headers = ['MODELO', 'VALOR (ate 35%)', 'VALOR CHEIO']
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        if i == 0:
            align = WD_ALIGN_PARAGRAPH.LEFT
        else:
            align = WD_ALIGN_PARAGRAPH.RIGHT
        style_cell(cell, h, size=10, bold=True,
                   color=WHITE, bg=CYBER_BLUE, align=align)
        set_cell_borders(cell, color=CYBER_BLUE, size='4')

    # Data rows
    for i, (model, loj, full) in enumerate(models):
        bg = WHITE if i % 2 == 0 else LIGHT_BG
        row = t.rows[i + 1]
        cells = row.cells

        # Modelo (col 0) — left align
        style_cell(cells[0], model, size=9.5, color=GRAY_700,
                   align=WD_ALIGN_PARAGRAPH.LEFT, bg=bg)
        set_cell_borders(cells[0], color=GRAY_200, size='4')

        # Valor 35% (col 1) — right align, cyber blue, bold
        style_cell(cells[1], brl(loj), size=10, bold=True,
                   color=CYBER_BLUE, align=WD_ALIGN_PARAGRAPH.RIGHT, bg=bg)
        set_cell_borders(cells[1], color=GRAY_200, size='4')

        # Valor cheio (col 2) — right align, gray
        style_cell(cells[2], brl(full), size=10, color=GRAY_500,
                   align=WD_ALIGN_PARAGRAPH.RIGHT, bg=bg)
        set_cell_borders(cells[2], color=GRAY_200, size='4')

    # Espaco entre blocos
    add_paragraph(doc, '', size=4, space_after=2)


# ============================================================================
# Build
# ============================================================================

def build_doc():
    doc = Document()

    # Margins da pagina
    section = doc.sections[0]
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)

    # ============================================================
    # HEADER
    # ============================================================
    add_paragraph(doc, 'CATALOGO 2026', size=10, bold=True,
                  color=CYBER_BLUE, space_after=2)
    add_paragraph(doc, 'Tabela de Precos - Atacado', size=20, bold=True,
                  color=NAVY_950, space_after=6)

    add_paragraph(
        doc,
        'Valor Cheio = valor medio de troca completa da tela em assistencia '
        'tecnica (peca + mao de obra). Valor (ate 35%) = quanto o lojista '
        'credenciado paga a Cyber. Tabela de referencia para cotacao. '
        'Modelos fora da lista: cotacao via WhatsApp.',
        size=10, color=GRAY_700, space_after=8,
    )

    # ============================================================
    # TABELAS POR MARCA
    # ============================================================
    for brand, models in GROUPS:
        add_brand_table(doc, brand, models)

    # ============================================================
    # CTA FINAL
    # ============================================================
    add_paragraph(doc, '', size=8, space_after=4)

    # Box CTA via tabela 1x1 com fundo
    cta = doc.add_table(rows=1, cols=1)
    cta.alignment = WD_TABLE_ALIGNMENT.CENTER
    cta.autofit = False
    cell = cta.rows[0].cells[0]
    cell.width = Cm(17)
    set_cell_bg(cell, LIGHT_BG)
    set_cell_borders(cell, color=CYBER_BLUE, size='4')

    cell.text = ''
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run('QUER CREDENCIAMENTO?')
    r1.font.size = Pt(11)
    r1.bold = True
    r1.font.color.rgb = RGBColor.from_string(CYBER_BLUE)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(
        'Acesse telas.cyberinformatica.tech ou fale no WhatsApp '
        '(11) 95436-9269. Analise de CNPJ em ate 24h uteis. '
        'Modelos fora da lista recebem cotacao personalizada.'
    )
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = RGBColor.from_string(GRAY_700)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run(
        'Valores em reais (R$) por unidade processada. Tabela valida para 2026 '
        '- atualizada periodicamente. Sujeito a analise cadastral.'
    )
    r3.font.size = Pt(7.5)
    r3.italic = True
    r3.font.color.rgb = RGBColor.from_string(GRAY_500)

    return doc


def main():
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc = build_doc()
    doc.save(OUTPUT)
    size = os.path.getsize(OUTPUT)
    print(f'OK  -> {OUTPUT} ({size:,} bytes, {TOTAL_MODELS} modelos)')


if __name__ == '__main__':
    main()